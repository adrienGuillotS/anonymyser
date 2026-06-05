import sys
import time

from docx import Document
from presidio_analyzer import AnalyzerEngine
from presidio_analyzer.nlp_engine import NlpEngineProvider
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig


def build_analyzer() -> AnalyzerEngine:
    configuration = {
        "nlp_engine_name": "spacy",
        "models": [{"lang_code": "fr", "model_name": "fr_core_news_lg"}],
    }
    provider = NlpEngineProvider(nlp_configuration=configuration)
    nlp_engine = provider.create_engine()
    return AnalyzerEngine(nlp_engine=nlp_engine, supported_languages=["fr"])


def build_operators() -> dict:
    return {
        "PERSON": OperatorConfig("replace", {"new_value": "<NAME>"}),
        "PHONE_NUMBER": OperatorConfig("replace", {"new_value": "<PHONE>"}),
        "EMAIL_ADDRESS": OperatorConfig("replace", {"new_value": "<EMAIL>"}),
        "LOCATION": OperatorConfig("replace", {"new_value": "<LOCATION>"}),
        "DATE_TIME": OperatorConfig("replace", {"new_value": "<DATE>"}),
        "DEFAULT": OperatorConfig("replace", {"new_value": "<REDACTED>"}),
    }


def anonymize_text(text: str, analyzer: AnalyzerEngine, anonymizer: AnonymizerEngine) -> str:
    if not (text or "").strip():
        return text

    results = analyzer.analyze(text=text, language="fr")
    anonymized = anonymizer.anonymize(
        text=text,
        analyzer_results=results,
        operators=build_operators(),
    )
    return anonymized.text


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def anonymize_docx(input_docx: str, output_docx: str) -> None:
    analyzer = build_analyzer()
    anonymizer = AnonymizerEngine()

    doc = Document(input_docx)

    for paragraph in iter_all_paragraphs(doc):
        original_text = paragraph.text
        if not (original_text or "").strip():
            continue

        new_text = anonymize_text(original_text, analyzer, anonymizer)
        if new_text == original_text:
            continue

        # Replace the whole paragraph text.
        # Note: this resets run-level formatting within the paragraph.
        paragraph.text = new_text

    doc.save(output_docx)


def main(argv: list[str]) -> int:
    if len(argv) < 3:
        print("Usage: python run_docx_replace.py input.docx output.docx", file=sys.stderr)
        return 2

    input_docx = argv[1]
    output_docx = argv[2]

    start = time.time()
    anonymize_docx(input_docx, output_docx)
    end = time.time()

    print(f"DOCX anonymisé généré: {output_docx}")
    print(f"temps {end - start}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
